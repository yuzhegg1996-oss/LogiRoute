# -*- coding: utf-8 -*-
"""
DOCX文档解析器

解析Word文档(.docx)的功能
移植自RAGFlow框架
集成时间: 2025年7月30日-8月2日
"""

import logging
from typing import List, Tuple, Optional
from io import BytesIO

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

class RAGFlowDocxParser:
    """
    DOCX文档解析器
    """
    
    def __init__(self):
        """初始化DOCX解析器"""
        if not HAS_DOCX:
            logger.warning("未安装python-docx库，请安装python-docx")
    
    def __call__(self, fnm: str, binary: Optional[bytes] = None) -> List[Tuple[str, str]]:
        """
        解析DOCX文档
        
        Args:
            fnm: 文件名
            binary: 二进制内容
            
        Returns:
            List[Tuple[str, str]]: 解析结果
        """
        if not HAS_DOCX:
            raise ImportError("需要安装python-docx库来解析DOCX文档")
        
        try:
            if binary:
                doc = Document(BytesIO(binary))
            else:
                doc = Document(fnm)
            
            results = []
            
            # 提取段落文本
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    # 判断是否为标题
                    style_name = paragraph.style.name.lower() if paragraph.style else ""
                    if 'heading' in style_name or 'title' in style_name:
                        results.append((text, f"heading_{i}"))
                    else:
                        results.append((text, f"paragraph_{i}"))
            
            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                table_content = self._extract_table_content(table)
                if table_content:
                    results.append((table_content, f"table_{table_idx}"))
            
            return results if results else [("无法提取文档内容", "error")]
            
        except Exception as e:
            logger.error(f"DOCX解析失败: {e}")
            return [(f"DOCX解析失败: {str(e)}", "error")]
    
    def _extract_table_content(self, table) -> str:
        """
        提取表格内容
        
        Args:
            table: 表格对象
            
        Returns:
            str: 表格文本内容
        """
        try:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(cell_text if cell_text else "")
                if any(cells):  # 至少有一个非空单元格
                    rows.append(" | ".join(cells))
            
            return "\n".join(rows) if rows else ""
            
        except Exception as e:
            logger.warning(f"提取表格内容时出错: {e}")
            return ""
    
    def _get_paragraph_style(self, paragraph) -> str:
        """
        获取段落样式类型
        
        Args:
            paragraph: 段落对象
            
        Returns:
            str: 样式类型
        """
        try:
            if paragraph.style:
                style_name = paragraph.style.name.lower()
                if 'heading' in style_name:
                    return 'heading'
                elif 'title' in style_name:
                    return 'title'
                elif 'caption' in style_name:
                    return 'caption'
            return 'paragraph'
        except:
            return 'paragraph'
    
    def extract_images_info(self, fnm: str, binary: Optional[bytes] = None) -> List[dict]:
        """
        提取文档中的图片信息
        
        Args:
            fnm: 文件名
            binary: 二进制内容
            
        Returns:
            List[dict]: 图片信息列表
        """
        if not HAS_DOCX:
            return []
        
        try:
            if binary:
                doc = Document(BytesIO(binary))
            else:
                doc = Document(fnm)
            
            images_info = []
            
            # 遍历所有关系，查找图片
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images_info.append({
                        'filename': rel.target_ref.split('/')[-1],
                        'content_type': rel.target_part.content_type if hasattr(rel, 'target_part') else 'unknown'
                    })
            
            return images_info
            
        except Exception as e:
            logger.warning(f"提取图片信息时出错: {e}")
            return []
if __name__ == "__main__":
    # docx_path = "C:\\Users\\woshi\\OneDrive\\桌面\\教研论文\\基于三阶映射模型的虚拟现实技术课程思政教学探索与实践.docx"
    docx_path = "C:\\Users\\woshi\\Downloads\\基于OpenCV的车辆检测系统设计与实现毕业论文 (1).docx"
    with open(docx_path, 'rb') as f:
        file_content = f.read()
    file_name=docx_path.split('\\')[-1]
    parser = RAGFlowDocxParser()
    images_info = parser.extract_images_info(file_name, binary=file_content)
    print(images_info)
    
    # 提取文档中的标题和正文
    print("\n开始提取文档内容...")
    doc_content = parser(file_name, binary=file_content)
    
    # 创建标题与段落内容的映射字典
    content_dict = {}
    current_heading = None
    current_paragraphs = []
    
    # 遍历解析结果，构建标题与段落的映射关系
    for content, content_type in doc_content:
        if content_type.startswith("heading_"):
            # 如果遇到新的标题，将之前收集的段落保存到字典中
            if current_heading is not None:
                content_dict[current_heading] = current_paragraphs
            
            # 更新当前标题，并重置段落列表
            current_heading = content
            current_paragraphs = []
        elif content_type.startswith("paragraph_"):
            # 如果当前有标题，则将段落添加到当前标题下
            if current_heading is not None:
                current_paragraphs.append(content)
            # 如果还没有遇到标题，将段落添加到一个特殊键下（如"无标题段落"）
            else:
                if "无标题段落" not in content_dict:
                    content_dict["无标题段落"] = []
                content_dict["无标题段落"].append(content)
        elif content_type.startswith("table_"):
            # 表格内容可以单独处理或添加到当前标题下
            if current_heading is not None:
                current_paragraphs.append(f"[表格内容]: {content}")
            else:
                if "无标题段落" not in content_dict:
                    content_dict["无标题段落"] = []
                content_dict["无标题段落"].append(f"[表格内容]: {content}")
        elif content_type == "error":
            print(f"解析错误: {content}")
    
    # 保存最后一个标题下的段落
    if current_heading is not None:
        content_dict[current_heading] = current_paragraphs
    
    # 打印构建的字典结构
    print("\n文档内容结构:")
    print("-" * 30)
    for heading, paragraphs in content_dict.items():
        print(f"标题: {heading}")
        # for i, paragraph in enumerate(paragraphs, 1):
        #     print(f"  段落{i}: {paragraph}")
        # print()